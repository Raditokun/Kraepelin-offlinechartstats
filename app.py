import tkinter as tk
from tkinter import messagebox, filedialog
import math
import os
import tempfile
import traceback
from datetime import datetime
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Inches, Pt

hasil_tes = {}

# Absolute path in the OS temp dir so it works regardless of the exe's working directory
TEMP_GRAFIK = os.path.join(tempfile.gettempdir(), "temp_grafik.png")

def penentuan_kategori(prestasi, tempo, konsistensi, ketahanan):
    kat_prestasi = "Kurang" if prestasi < 500 else "Cukup" if prestasi <= 700 else "Baik"
    kat_tempo = "Rendah" if tempo < 15 else "Sedang" if tempo <= 22 else "Cepat"
    kat_konsistensi = "Sedang" if konsistensi < 3.0 else "Buruk" # Semakin kecil nilai, semakin konsisten
    kat_ketahanan = "Sedang" if -0.5 <= ketahanan <= 0.5 else "Baik" if ketahanan > 0.5 else "Buruk"
    
    return kat_prestasi, kat_tempo, kat_konsistensi, kat_ketahanan

def hitung_statistik():
    input_text = entry_data.get()
    input_ketelitian = entry_ketelitian.get()
    
    try:
        data = [int(x) for x in input_text.strip().split()]
        ketelitian = int(input_ketelitian) if input_ketelitian else 0

        prestasi = sum(data)
        tempo = prestasi / len(data)
        variance = sum((x - tempo) ** 2 for x in data) / len(data)
        konsistensi = math.sqrt(variance)
        
        n = len(data)
        if n < 2:
            messagebox.showerror("Error", "Data harus berisi setidaknya 2 nilai untuk menghitung ketahanan.")
            return
        sum_x = sum(range(1, n + 1))
        sum_y = sum(data)
        sum_xy = sum((i + 1) * data[i] for i in range(n))
        sum_xx = sum((i + 1) ** 2 for i in range(n))
        ketahanan = ((n * sum_xy) - (sum_x * sum_y)) / ((n * sum_xx) - (sum_x ** 2))

        
        kat_pres, kat_temp, kat_kon, kat_ket = penentuan_kategori(prestasi, tempo, konsistensi, ketahanan)

       
        global hasil_tes
        hasil_tes = {
            "prestasi": prestasi, "tempo": tempo, "konsistensi": konsistensi,
            "ketahanan": ketahanan, "ketelitian": ketelitian,
            "min": min(data), "max": max(data),
            "kat_prestasi": kat_pres, "kat_tempo": kat_temp,
            "kat_konsistensi": kat_kon, "kat_ketahanan": kat_ket,
            "nama": entry_nama.get().strip(),
            "posisi": entry_posisi.get().strip(),
            "perusahaan": entry_perusahaan.get().strip(),
            "tanggal": entry_tanggal.get().strip()
        }

        
        lbl_hasil.config(text=(
            f"Prestasi: {prestasi} ({kat_pres}) | Tempo: {tempo:.2f} ({kat_temp})\n"
            f"Konsistensi: {konsistensi:.4f} ({kat_kon}) | Ketahanan: {ketahanan:.5f} ({kat_ket})\n"
            f"Ketelitian: {ketelitian} | Min: {min(data)} | Max: {max(data)}"
        ))

        
        #Render Grafik
        ax.clear()
        ax.plot(range(1, n + 1), data, marker='o', color='#c0392b', linestyle='-', markersize=5)
        ax.set_title(f"Prestasi Kerja Tiap 30 Detik (Total: {n} Sesi)", fontsize=10)
        ax.set_xlim(1, 45)
        ax.set_xticks(range(1, 46))
        ax.tick_params(axis='x', labelsize=7) 
        ax.set_ylim(0, 60)               
        ax.set_yticks(range(0, 61, 5))  
        ax.grid(True, linestyle='--', alpha=0.5)
        for i, val in enumerate(data):
            ax.annotate(str(val), (i + 1, val), textcoords="offset points", xytext=(0,5), ha='center', fontsize=7)
        
        
        fig.savefig(TEMP_GRAFIK, dpi=150, bbox_inches='tight')
        canvas.draw()
        
        
        btn_word.config(state=tk.NORMAL)

    except ValueError:
        messagebox.showerror("Error", "Pastikan input grafik berisi angka dipisahkan spasi, dan ketelitian berisi angka!")
    except Exception:
        messagebox.showerror("Error", f"Gagal generate laporan:\n\n{traceback.format_exc()}")

def ekspor_word():
    if not hasil_tes:
        messagebox.showwarning("Peringatan", "Silakan generate laporan terlebih dahulu!")
        return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word documents", "*.docx")],
        title="Simpan Laporan Tes"
    )

    if not file_path:
        return

    try:
        doc = Document()

        judul = doc.add_heading("Laporan Hasil Tes Kraepelin", level=1)

        # Grafik (tersimpan silent sebagai temp_grafik.png, disisipkan ke dokumen)
        doc.add_picture(TEMP_GRAFIK, width=Inches(6.5))

        # Data Peserta
        doc.add_heading("Data Peserta", level=2)
        for label, key in [("Nama", "nama"), ("Posisi", "posisi"),
                           ("Perusahaan", "perusahaan"), ("Tanggal Ujian", "tanggal")]:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(hasil_tes[key])

        # Tabel Hasil Statistik
        doc.add_heading("Hasil Statistik", level=2)
        data_tabel = [
            ("Jumlah (Prestasi Kerja)", str(hasil_tes['prestasi']), hasil_tes['kat_prestasi']),
            ("Rata-rata (Tempo Kerja)", f"{hasil_tes['tempo']:.2f}", hasil_tes['kat_tempo']),
            ("Ketelitian", str(hasil_tes['ketelitian']), "Sangat Buruk" if hasil_tes['ketelitian'] > 20 else "Baik"),
            ("Konsistensi", f"{hasil_tes['konsistensi']:.4f}", hasil_tes['kat_konsistensi']),
            ("Ketahanan", f"{hasil_tes['ketahanan']:.5f}", hasil_tes['kat_ketahanan']),
            ("Nilai Minimal", str(hasil_tes['min']), ""),
            ("Nilai Maksimal", str(hasil_tes['max']), "")
        ]

        tabel = doc.add_table(rows=1, cols=3)
        tabel.style = "Table Grid"
        header = tabel.rows[0].cells
        for i, judul_kolom in enumerate(("Aspek", "Nilai", "Kategori")):
            header[i].text = judul_kolom
            header[i].paragraphs[0].runs[0].bold = True

        for aspek, nilai, kategori in data_tabel:
            baris = tabel.add_row().cells
            baris[0].text = aspek
            baris[1].text = nilai
            baris[2].text = kategori

        doc.save(file_path)

        if os.path.exists(TEMP_GRAFIK):
            os.remove(TEMP_GRAFIK)

        messagebox.showinfo("Sukses", f"Laporan Word berhasil disimpan di:\n{file_path}")

    except Exception:
        messagebox.showerror("Error", f"Gagal membuat Word:\n\n{traceback.format_exc()}")

#ui
root = tk.Tk()
root.title("Evaluator Kraepelin Internal HR - Word Ready")
root.geometry("850x650")

frame_top = tk.Frame(root, pady=10)
frame_top.pack(fill=tk.X)

#Baris 0 - Identitas Peserta
frame_identitas = tk.Frame(frame_top)
frame_identitas.pack(pady=5)

tk.Label(frame_identitas, text="Nama:", font=("Arial", 9, "bold")).grid(row=0, column=0, sticky="e", padx=5, pady=2)
entry_nama = tk.Entry(frame_identitas, width=30)
entry_nama.grid(row=0, column=1, padx=5, pady=2)

tk.Label(frame_identitas, text="Posisi:", font=("Arial", 9, "bold")).grid(row=0, column=2, sticky="e", padx=5, pady=2)
entry_posisi = tk.Entry(frame_identitas, width=30)
entry_posisi.grid(row=0, column=3, padx=5, pady=2)

tk.Label(frame_identitas, text="Perusahaan:", font=("Arial", 9, "bold")).grid(row=1, column=0, sticky="e", padx=5, pady=2)
entry_perusahaan = tk.Entry(frame_identitas, width=30)
entry_perusahaan.grid(row=1, column=1, padx=5, pady=2)

tk.Label(frame_identitas, text="Tanggal Ujian:", font=("Arial", 9, "bold")).grid(row=1, column=2, sticky="e", padx=5, pady=2)
entry_tanggal = tk.Entry(frame_identitas, width=30)
entry_tanggal.grid(row=1, column=3, padx=5, pady=2)

#Baris 1
tk.Label(frame_top, text="Masukkan 50 Nilai Puncak (Spasi):", font=("Arial", 9, "bold")).pack()
entry_data = tk.Entry(frame_top, width=100)
entry_data.pack(pady=3)
entry_data.insert(0, "6 11 6 8 8 9 9 8 6 6 4 5 4 6 7 4 8 7 10 7 7 5 6 8 5 11 8 9 10 8 7 10 5 10 7 7 8 5 5 11 11 8 6 6 7 8 10 10 8 7")

#Baris 2
tk.Label(frame_top, text="Jumlah Kesalahan/Terlewat (Ketelitian):", font=("Arial", 9, "bold")).pack()
entry_ketelitian = tk.Entry(frame_top, width=20)
entry_ketelitian.pack(pady=3)
entry_ketelitian.insert(0, "26")
#Baris 3
frame_btn = tk.Frame(frame_top)
frame_btn.pack(pady=5)
tk.Button(frame_btn, text="Generate Laporan", command=hitung_statistik, bg="#3498db", fg="white", width=20).pack(side=tk.LEFT, padx=10)
btn_word = tk.Button(frame_btn, text="Simpan Laporan Word", command=ekspor_word, bg="#2ecc71", fg="white", width=20, state=tk.DISABLED)
btn_word.pack(side=tk.LEFT, padx=10)
lbl_hasil = tk.Label(root, text="Hasil statistik akan muncul di sini.", font=("Consolas", 10), bg="#f0f0f0", relief="sunken", padx=10, pady=10)
lbl_hasil.pack(fill=tk.X, padx=20, pady=5)
fig = Figure(figsize=(8, 3.5), dpi=100)
ax = fig.add_subplot(111)
canvas = FigureCanvasTkAgg(fig, master=root)
canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

root.mainloop()